from __future__ import annotations

import io
import json
import os
import re
from csv import reader as csv_reader
from pathlib import Path
from typing import Any

import requests
import streamlit as st
from docx import Document
from openai import OpenAI
from openai import APIStatusError
from openpyxl import load_workbook
from pypdf import PdfReader


BASE_DIR = Path(__file__).resolve().parent
ANZSCO_DIR = BASE_DIR / "anzsco"
SIFA_DIR = BASE_DIR / "sifa"
ADMIN_PROMPTS_DIR = BASE_DIR / "admin_prompts"
ADMIN_PROMPT_CONFIG_PATH = ADMIN_PROMPTS_DIR / "config.json"
ADMIN_USERNAME = "qibaitintern"
ADMIN_PASSWORD = "qibaitintern"
PROMPT_KEYS = {
    "resume": "resume_default_prompt",
    "interview_questions": "interview_questions_default_prompt",
    "interview_analysis": "interview_analysis_default_prompt",
}
SUPPORTED_REFERENCE_EXTENSIONS = {".pdf", ".doc", ".docx", ".csv", ".xlsx", ".txt"}
MAX_REFERENCE_CHARS_PER_FILE = 6000
MAX_REFERENCE_CHARS_PER_FOLDER = 24000
OPENROUTER_BASE_URL = os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")
DEFAULT_OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1")
DEFAULT_SCORE_ROWS = 5


def ensure_app_directories() -> None:
    ANZSCO_DIR.mkdir(exist_ok=True)
    SIFA_DIR.mkdir(exist_ok=True)
    ADMIN_PROMPTS_DIR.mkdir(exist_ok=True)


def extract_text_from_doc_bytes(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore").replace("\x00", " ").strip()


def extract_text_from_xlsx_bytes(data: bytes) -> str:
    workbook = load_workbook(io.BytesIO(data), data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"[Sheet] {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell) for cell in row if cell is not None and str(cell).strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts).strip()


def extract_text_from_csv_bytes(data: bytes) -> str:
    text = data.decode("utf-8", errors="ignore")
    rows = []
    for row in csv_reader(io.StringIO(text)):
        values = [cell.strip() for cell in row if cell and cell.strip()]
        if values:
            rows.append(" | ".join(values))
    return "\n".join(rows).strip()


def extract_text_from_upload(uploaded_file) -> str:
    if uploaded_file is None:
        return ""

    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()

    if name.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()

    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore").strip()

    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(page.strip() for page in pages if page.strip()).strip()

    if name.endswith(".csv"):
        return extract_text_from_csv_bytes(data)

    if name.endswith(".xlsx"):
        return extract_text_from_xlsx_bytes(data)

    if name.endswith(".doc"):
        return extract_text_from_doc_bytes(data)

    raise ValueError("Unsupported file type. Please upload a DOC, DOCX, PDF, TXT, CSV, or XLSX file.")


def extract_text_from_path(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    data = file_path.read_bytes()

    if suffix == ".docx":
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()

    if suffix == ".txt":
        return data.decode("utf-8", errors="ignore").strip()

    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(page.strip() for page in pages if page.strip()).strip()

    if suffix == ".csv":
        return extract_text_from_csv_bytes(data)

    if suffix == ".xlsx":
        return extract_text_from_xlsx_bytes(data)

    if suffix == ".doc":
        return extract_text_from_doc_bytes(data)

    return ""


def get_openai_client() -> OpenAI:
    api_key = (
        st.session_state.get("api_key_input")
        or os.getenv("OPENROUTER_API_KEY")
        or os.getenv("OPENAI_API_KEY")
    )
    if not api_key:
        raise RuntimeError(
            "Missing API key. Set OPENROUTER_API_KEY for OpenRouter or OPENAI_API_KEY for OpenAI."
        )

    using_openrouter = bool(st.session_state.get("api_key_input") or os.getenv("OPENROUTER_API_KEY"))
    base_url = OPENROUTER_BASE_URL if using_openrouter else DEFAULT_OPENAI_BASE_URL
    return OpenAI(api_key=api_key, base_url=base_url)


def get_default_model() -> str:
    return os.getenv("OPENAI_MODEL") or os.getenv("OPENROUTER_MODEL") or "openai/gpt-5"


def format_price_per_million(value: str | None) -> str:
    if not value:
        return "$0/M"

    try:
        per_token = float(value)
    except (TypeError, ValueError):
        return value

    per_million = per_token * 1_000_000
    if per_million == 0:
        return "$0/M"
    if per_million >= 1:
        return f"${per_million:,.2f}/M"
    return f"${per_million:.4f}/M"


def format_context_length(value: int | None) -> str:
    if not value:
        return "N/A"
    return f"{value:,}"


@st.cache_data(show_spinner=False, ttl=900)
def fetch_openrouter_free_models(api_key: str) -> list[dict[str, Any]]:
    response = requests.get(
        f"{OPENROUTER_BASE_URL}/models",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Accept": "application/json",
        },
        timeout=20,
    )
    response.raise_for_status()

    data = response.json().get("data", [])
    free_models = []
    for item in data:
        model_id = item.get("id", "")
        if ":free" in model_id:
            pricing = item.get("pricing", {}) or {}
            context_length = item.get("context_length") or item.get("top_provider", {}).get("context_length")
            free_models.append(
                {
                    "id": model_id,
                    "name": item.get("name") or model_id,
                    "context_length": context_length,
                    "pricing": pricing,
                    "label": (
                        f"{model_id} | "
                        f"Input {format_price_per_million(pricing.get('prompt'))} | "
                        f"Output {format_price_per_million(pricing.get('completion'))} | "
                        f"Context {format_context_length(context_length)}"
                    ),
                }
            )

    unique_models = {model["id"]: model for model in free_models}
    return [unique_models[key] for key in sorted(unique_models.keys())]


def get_available_models() -> tuple[list[dict[str, Any]], str | None]:
    openrouter_api_key = st.session_state.get("api_key_input") or os.getenv("OPENROUTER_API_KEY")
    if not openrouter_api_key:
        configured_model = get_default_model()
        return [
            {
                "id": configured_model,
                "name": configured_model,
                "context_length": None,
                "pricing": {},
                "label": f"{configured_model} | Input N/A | Output N/A | Context N/A",
            }
        ], None

    try:
        free_models = fetch_openrouter_free_models(openrouter_api_key)
        if free_models:
            return free_models, None
        fallback_model = get_default_model()
        return [
            {
                "id": fallback_model,
                "name": fallback_model,
                "context_length": None,
                "pricing": {},
                "label": f"{fallback_model} | Input N/A | Output N/A | Context N/A",
            }
        ], None
    except Exception as exc:
        fallback_model = get_default_model()
        return [
            {
                "id": fallback_model,
                "name": fallback_model,
                "context_length": None,
                "pricing": {},
                "label": f"{fallback_model} | Input N/A | Output N/A | Context N/A",
            }
        ], f"Could not fetch OpenRouter models dynamically: {exc}"


def get_model_by_id(models: list[dict[str, Any]], model_id: str) -> dict[str, Any]:
    for model in models:
        if model["id"] == model_id:
            return model
    return models[0]


def get_model_retry_order(models: list[dict[str, Any]], preferred_model_id: str) -> list[str]:
    model_ids = [model["id"] for model in models]
    if preferred_model_id not in model_ids:
        return model_ids
    return [preferred_model_id] + [model_id for model_id in model_ids if model_id != preferred_model_id]


def extract_retry_after_seconds(exc: APIStatusError) -> float:
    response = getattr(exc, "response", None)
    if response is None:
        return 0

    headers = getattr(response, "headers", {}) or {}
    retry_after = headers.get("Retry-After")
    if retry_after:
        try:
            return float(retry_after)
        except ValueError:
            return 0
    return 0


def call_llm(system_prompt: str, payload: dict[str, Any], model: str, available_models: list[dict[str, Any]]) -> str:
    client = get_openai_client()
    last_error = None

    for candidate_model in get_model_retry_order(available_models, model):
        try:
            st.session_state["active_request_model"] = candidate_model
            response = client.chat.completions.create(
                model=candidate_model,
                messages=[
                    {
                        "role": "system",
                        "content": system_prompt,
                    },
                    {
                        "role": "user",
                        "content": json.dumps(payload, ensure_ascii=False, indent=2),
                    },
                ],
            )
            st.session_state["selected_model"] = candidate_model
            return (response.choices[0].message.content or "").strip()
        except APIStatusError as exc:
            last_error = exc
            if exc.status_code == 429:
                retry_after_seconds = extract_retry_after_seconds(exc)
                if retry_after_seconds > 0 and candidate_model == model:
                    st.warning(
                        f"Model `{candidate_model}` is rate-limited. Retrying with another free model now."
                    )
                continue
            raise

    if last_error is not None:
        raise RuntimeError(
            "All available free models are temporarily rate-limited. "
            "Please retry shortly or choose a different provider/account."
        ) from last_error

    raise RuntimeError("No available model could be used for this request.")


def create_word_report(title: str, sections: list[tuple[str, str]]) -> bytes:
    doc = Document()
    doc.add_heading(title, 0)

    for heading, content in sections:
        doc.add_heading(heading, level=1)
        for block in split_blocks(content):
            doc.add_paragraph(block)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def split_blocks(text: str) -> list[str]:
    blocks = [block.strip() for block in text.split("\n") if block.strip()]
    return blocks or ["No content generated."]


def read_admin_prompt_config() -> dict[str, str]:
    if not ADMIN_PROMPT_CONFIG_PATH.exists():
        return {}
    try:
        return json.loads(ADMIN_PROMPT_CONFIG_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}


def write_admin_prompt_config(config: dict[str, str]) -> None:
    ADMIN_PROMPT_CONFIG_PATH.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")


def save_admin_prompt(uploaded_file, prompt_key: str) -> Path:
    config = read_admin_prompt_config()
    suffix = Path(uploaded_file.name).suffix.lower() or ".txt"
    target_path = ADMIN_PROMPTS_DIR / f"{prompt_key}{suffix}"

    for existing in ADMIN_PROMPTS_DIR.glob(f"{prompt_key}.*"):
        if existing.is_file():
            existing.unlink()

    target_path.write_bytes(uploaded_file.getvalue())
    config[prompt_key] = target_path.name
    write_admin_prompt_config(config)
    return target_path


def get_admin_prompt_text(prompt_key: str) -> str:
    config = read_admin_prompt_config()
    file_name = config.get(prompt_key)
    if not file_name:
        return ""

    file_path = ADMIN_PROMPTS_DIR / file_name
    if not file_path.exists():
        return ""

    return extract_text_from_path(file_path)


def resolve_prompt_text(uploaded_file, prompt_key: str) -> tuple[str, str]:
    if uploaded_file is not None:
        return extract_text_from_upload(uploaded_file), "uploaded"

    admin_prompt_text = get_admin_prompt_text(prompt_key)
    if admin_prompt_text:
        return admin_prompt_text, "admin_default"

    return "", "missing"


def load_reference_folder(folder_path: Path) -> str:
    collected_parts: list[str] = []
    current_chars = 0

    for file_path in sorted(folder_path.rglob("*")):
        if not file_path.is_file() or file_path.suffix.lower() not in SUPPORTED_REFERENCE_EXTENSIONS:
            continue

        try:
            extracted = extract_text_from_path(file_path).strip()
        except Exception:
            extracted = ""

        if not extracted:
            continue

        trimmed = extracted[:MAX_REFERENCE_CHARS_PER_FILE]
        block = f"[File] {file_path.name}\n{trimmed}"
        if current_chars + len(block) > MAX_REFERENCE_CHARS_PER_FOLDER:
            remaining = MAX_REFERENCE_CHARS_PER_FOLDER - current_chars
            if remaining <= 0:
                break
            block = block[:remaining]
        collected_parts.append(block)
        current_chars += len(block)
        if current_chars >= MAX_REFERENCE_CHARS_PER_FOLDER:
            break

    return "\n\n".join(collected_parts).strip()


def build_reference_payload() -> dict[str, str]:
    return {
        "anzsco_reference": load_reference_folder(ANZSCO_DIR),
        "sifa_reference": load_reference_folder(SIFA_DIR),
    }


RESUME_TAB_NAMES = [
    "JD Review",
    "Resume Review",
    "ATS Score",
    "JD Scorecard",
    "Resume Rewrite",
    "Strengthening",
    "Change Comparison",
    "Cover Letter",
    "Interview Preparation",
    "Iteration Tracking",
    "Readiness Analysis",
]


def parse_resume_report_sections(response_text: str) -> list[dict[str, str]]:
    marker_sections = parse_resume_report_sections_from_markers(response_text)
    if marker_sections is not None:
        return marker_sections

    cleaned_text = response_text.strip()
    fenced_match = re.search(r"```(?:json)?\s*(\{[\s\S]*\})\s*```", cleaned_text)
    if fenced_match:
        cleaned_text = fenced_match.group(1).strip()
    else:
        json_match = re.search(r"(\{[\s\S]*\})", cleaned_text)
        if json_match:
            cleaned_text = json_match.group(1).strip()

    try:
        parsed = json.loads(cleaned_text)
        sections = parsed.get("sections", [])
        normalized_sections = []
        for index, tab_name in enumerate(RESUME_TAB_NAMES):
            item = sections[index] if index < len(sections) else {}
            normalized_sections.append(
                {
                    "tab_name": tab_name,
                    "title": str(item.get("title") or tab_name),
                    "content": str(item.get("content") or "No content generated."),
                }
            )
        return normalized_sections
    except Exception:
        return [
            {
                "tab_name": RESUME_TAB_NAMES[0],
                "title": RESUME_TAB_NAMES[0],
                "content": response_text,
            }
        ] + [
            {
                "tab_name": tab_name,
                "title": tab_name,
                "content": "No content generated.",
            }
            for tab_name in RESUME_TAB_NAMES[1:]
        ]


def parse_resume_report_sections_from_markers(response_text: str) -> list[dict[str, str]] | None:
    lines = response_text.splitlines()
    extracted_sections: dict[str, list[str]] = {}
    current_tab: str | None = None

    for line in lines:
        stripped_line = line.strip()
        matched_tab = next(
            (
                tab_name
                for tab_name in RESUME_TAB_NAMES
                if stripped_line == f"[[{tab_name}]]"
            ),
            None,
        )
        if matched_tab:
            current_tab = matched_tab
            extracted_sections[current_tab] = []
            continue

        if current_tab is not None:
            extracted_sections[current_tab].append(line)

    if not extracted_sections:
        return None

    normalized_sections = []
    for tab_name in RESUME_TAB_NAMES:
        content = "\n".join(extracted_sections.get(tab_name, [])).strip() or "No content generated."
        normalized_sections.append(
            {
                "tab_name": tab_name,
                "title": tab_name,
                "content": content,
            }
        )
    return normalized_sections


def build_score_rows(rows: int) -> list[dict[str, Any]]:
    return [
        {
            "Question Description": "",
            "S": 0.0,
            "T/E": 0.0,
            "A": 0.0,
            "R/T": 0.0,
            "STAR Rating": 0.0,
            "C/S": 0.0,
            "A/E": 0.0,
            "R/A": 0.0,
            "Competency Rating": 0.0,
        }
        for _ in range(rows)
    ]


def recalculate_score_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    recalculated_rows = []
    for row in rows:
        s_score = float(row.get("S", 0.0) or 0.0)
        t_score = float(row.get("T/E", 0.0) or 0.0)
        a_score = float(row.get("A", 0.0) or 0.0)
        r_score = float(row.get("R/T", 0.0) or 0.0)
        c_score = float(row.get("C/S", 0.0) or 0.0)
        ae_score = float(row.get("A/E", 0.0) or 0.0)
        ra_score = float(row.get("R/A", 0.0) or 0.0)
        recalculated_rows.append(
            {
                "Question Description": row.get("Question Description", ""),
                "S": s_score,
                "T/E": t_score,
                "A": a_score,
                "R/T": r_score,
                "STAR Rating": round((s_score + t_score + a_score + r_score) / 4, 2),
                "C/S": c_score,
                "A/E": ae_score,
                "R/A": ra_score,
                "Competency Rating": round((c_score + ae_score + ra_score) / 3, 2),
            }
        )
    return recalculated_rows


def build_metric_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    star_ratings = [float(row.get("STAR Rating", 0.0) or 0.0) for row in rows]
    competency_ratings = [float(row.get("Competency Rating", 0.0) or 0.0) for row in rows]
    all_ratings = star_ratings + competency_ratings
    total_ratings = len(all_ratings)
    average_score = round(sum(all_ratings) / total_ratings, 2) if total_ratings else 0.0
    strong_count = sum(1 for rating in all_ratings if rating >= 0.75)
    weak_count = sum(1 for rating in all_ratings if rating < 0.5)

    return [
        {"Metric": "average score", "Value": average_score},
        {"Metric": "strong (>= 0.75)", "Value": strong_count},
        {"Metric": "weak (< 0.5)", "Value": weak_count},
    ]


def render_simple_table(rows: list[dict[str, Any]], headers: list[str]) -> None:
    header_cols = st.columns(len(headers))
    for index, header in enumerate(headers):
        header_cols[index].markdown(f"**{header}**")

    for row in rows:
        row_cols = st.columns(len(headers))
        for index, header in enumerate(headers):
            row_cols[index].write(row.get(header, ""))


def render_readonly_score(value: float) -> None:
    st.markdown(
        f"""
        <div style="
            background-color: #f3f4f6;
            border-radius: 0.75rem;
            padding: 0.72rem 1rem;
            color: #6b7280;
            text-align: left;
            font-size: 1.1rem;
            line-height: 1.2;
            min-height: 3rem;
            display: flex;
            align-items: center;
            box-sizing: border-box;
        ">{value:.2f}</div>
        """,
        unsafe_allow_html=True,
    )


def render_download(report_name: str, report_bytes: bytes, button_label: str) -> None:
    st.download_button(
        label=button_label,
        data=report_bytes,
        file_name=report_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )


def persist_report(report_key: str, file_name: str, title: str, body: str, sections: list[tuple[str, str]]) -> None:
    st.session_state[report_key] = {
        "body": body,
        "file_name": file_name,
        "bytes": create_word_report(title, sections),
    }


def render_persisted_report(report_key: str, section_title: str, button_label: str) -> None:
    report = st.session_state.get(report_key)
    if not report:
        return

    st.markdown(f"### {section_title}")
    st.markdown(report["body"])
    render_download(report["file_name"], report["bytes"], button_label)


def render_resume_report_tabs(report_key: str, section_title: str, button_label: str) -> None:
    report = st.session_state.get(report_key)
    if not report:
        return

    st.markdown(f"### {section_title}")
    tab_sections = report.get("tab_sections", [])
    if not tab_sections and report.get("raw_response"):
        tab_sections = parse_resume_report_sections(report["raw_response"])
        st.session_state[report_key]["tab_sections"] = tab_sections
    if not tab_sections:
        st.markdown(report["body"])
        render_download(report["file_name"], report["bytes"], button_label)
        return

    tabs = st.tabs([section["tab_name"] for section in tab_sections])
    for tab, section in zip(tabs, tab_sections):
        with tab:
            if section["tab_name"] == "JD Review":
                render_jd_review_content(section["content"])
            else:
                st.markdown(section["content"], unsafe_allow_html=True)

    render_download(report["file_name"], report["bytes"], button_label)


def render_jd_review_content(content: str) -> None:
    colored_content = colorize_mandatory_requirements_content(content)
    st.markdown(colored_content, unsafe_allow_html=True)


def colorize_mandatory_requirements_content(content: str) -> str:
    start_marker = "[[MMR_START]]"
    end_marker = "[[MMR_END]]"

    start_index = content.find(start_marker)
    end_index = content.find(end_marker)
    if start_index == -1 or end_index == -1 or end_index <= start_index:
        return content

    prefix = content[:start_index]
    highlighted_body = content[start_index + len(start_marker) : end_index].strip()
    suffix = content[end_index + len(end_marker) :]
    highlighted_body = strip_mandatory_requirements_heading(highlighted_body)

    if highlighted_body:
        highlighted_body = (
            f'<div style="color: #dc2626; margin-bottom: 1rem;">{highlighted_body}</div>'
        )

    return f"{prefix}{highlighted_body}{suffix}"


def strip_mandatory_requirements_heading(content: str) -> str:
    pattern = re.compile(
        r"^\s*(?:\*\*)?Mandatory Minimum Requirements(?:\*\*)?\s*[-:–—]?\s*",
        re.IGNORECASE,
    )
    return re.sub(pattern, "", content, count=1).strip()


def render_resume_module() -> None:
    st.subheader("Resume")
    job_description = st.text_area("Job Description", height=220, key="resume_jd")
    resume_file = st.file_uploader(
        "Upload Resume",
        type=["doc", "docx", "pdf", "txt"],
        key="resume_file",
    )
    prompt_file = st.file_uploader(
        "Upload Analysis Prompt",
        type=["doc", "docx", "pdf", "txt"],
        key="resume_prompt",
    )
    prompt_notes = st.text_area(
        "Prompt Notes",
        height=220,
        key="resume_prompt_notes",
    )

    if st.button("Generate Resume Report", use_container_width=True):
        try:
            resume_text = extract_text_from_upload(resume_file)
            prompt_text, prompt_source = resolve_prompt_text(prompt_file, PROMPT_KEYS["resume"])

            if not job_description.strip():
                st.error("Please provide a job description.")
                return
            if not resume_text:
                st.error("Please upload a resume file with readable content.")
                return
            if not prompt_text:
                st.error("Please upload a prompt document or ask an admin to configure the default resume prompt.")
                return
            if prompt_source == "admin_default":
                st.info("Using the admin default resume prompt.")

            with st.spinner("Generating resume report..."):
                report_body = call_llm(
                    system_prompt=(
                        "You are an expert resume reviewer. Follow the uploaded prompt carefully and "
                        "produce a polished report in clear professional English. "
                        "Do not return JSON. Do not return code fences. Do not wrap the answer in XML or HTML containers. "
                        "Always return plain readable text split into exactly 11 sections using these exact markers on their own lines: "
                        "[[JD Review]], [[Resume Review]], [[ATS Score]], [[JD Scorecard]], [[Resume Rewrite]], "
                        "[[Strengthening]], [[Change Comparison]], [[Cover Letter]], [[Interview Preparation]], "
                        "[[Iteration Tracking]], [[Readiness Analysis]]. "
                        "If the user prompt includes [[MMR_START]] and [[MMR_END]] inside JD Review, preserve them exactly "
                        "around that body content only as invisible delimiters for UI styling. Do not let these markers "
                        "change the analysis, conclusions, scoring, level of detail, or wording choices. "
                        "After each marker, write only the content for that section until the next marker. "
                        "Keep the content human-readable and well formatted in markdown-style plain text."
                    ),
                    model=st.session_state["selected_model"],
                    available_models=st.session_state["available_models"],
                    payload={
                        "task": "resume_analysis",
                        "job_description": job_description,
                        "resume_text": resume_text,
                        "custom_prompt": prompt_text,
                        "prompt_notes": prompt_notes,
                        **build_reference_payload(),
                    },
                )

            tab_sections = parse_resume_report_sections(report_body)
            persist_report(
                "resume_report",
                "resume_report.docx",
                "Resume Analysis Report",
                "\n\n".join(
                    [f"{section['title']}\n\n{section['content']}" for section in tab_sections]
                ),
                [(section["title"], section["content"]) for section in tab_sections]
                + [("Job Description", job_description)],
            )
            st.session_state["resume_report"]["tab_sections"] = tab_sections
            st.session_state["resume_report"]["raw_response"] = report_body

            st.session_state["resume_report_text"] = st.session_state["resume_report"]["body"]
            st.success("Resume report generated.")
        except Exception as exc:
            st.error(str(exc))

    render_resume_report_tabs("resume_report", "Resume Report", "Download Resume Report")


def render_question_simulation() -> None:
    st.markdown("### Question Simulation")
    job_description = st.text_area("Job Description", height=220, key="interview_jd")
    resume_file = st.file_uploader(
        "Upload Resume",
        type=["doc", "docx", "pdf", "txt"],
        key="interview_resume_file",
    )
    prompt_file = st.file_uploader(
        "Upload Interview Prompt",
        type=["doc", "docx", "pdf", "txt"],
        key="interview_prompt",
    )

    if st.button("Generate Interview Questions", use_container_width=True):
        try:
            resume_text = extract_text_from_upload(resume_file)
            prompt_text, prompt_source = resolve_prompt_text(prompt_file, PROMPT_KEYS["interview_questions"])

            if not job_description.strip():
                st.error("Please provide a job description.")
                return
            if not resume_text:
                st.error("Please upload a resume file with readable content.")
                return
            if not prompt_text:
                st.error("Please upload a prompt document or ask an admin to configure the default interview question prompt.")
                return
            if prompt_source == "admin_default":
                st.info("Using the admin default interview question prompt.")

            with st.spinner("Generating interview question report..."):
                question_report = call_llm(
                    system_prompt=(
                        "You are an interview coach. Follow the uploaded prompt carefully and generate "
                        "likely interview questions in polished professional English."
                    ),
                    model=st.session_state["selected_model"],
                    available_models=st.session_state["available_models"],
                    payload={
                        "task": "interview_question_simulation",
                        "job_description": job_description,
                        "resume_text": resume_text,
                        "custom_prompt": prompt_text,
                        **build_reference_payload(),
                    },
                )

            persist_report(
                "interview_question_report",
                "interview_question_report.docx",
                "Interview Question Simulation Report",
                question_report,
                [
                    ("Simulated Questions", question_report),
                    ("Job Description", job_description),
                ],
            )

            st.session_state["question_report_text"] = question_report
            st.success("Interview questions generated.")
        except Exception as exc:
            st.error(str(exc))

    render_persisted_report(
        "interview_question_report",
        "Interview Question Report",
        "Download Interview Question Report",
    )


def render_scoring_section() -> list[dict[str, Any]]:
    st.markdown("### Scoring")

    if "score_row_count" not in st.session_state:
        st.session_state["score_row_count"] = DEFAULT_SCORE_ROWS
    if "score_rows" not in st.session_state:
        st.session_state["score_rows"] = build_score_rows(DEFAULT_SCORE_ROWS)

    row_count = st.number_input(
        "Number of Rows",
        min_value=1,
        max_value=50,
        value=int(st.session_state["score_row_count"]),
        step=1,
    )

    if row_count != st.session_state["score_row_count"]:
        st.session_state["score_row_count"] = row_count
        st.session_state["score_rows"] = build_score_rows(int(row_count))

    headers = ["Question Description", "S", "T/E", "A", "R/T", "Rating", "C/S", "A/E", "R/A", "Rating"]
    header_cols = st.columns([3, 1, 1, 1, 1, 1, 1, 1, 1, 1])
    for index, header in enumerate(headers):
        header_cols[index].markdown(f"**{header}**")

    updated_rows = []
    for index in range(int(row_count)):
        current_row = st.session_state["score_rows"][index]
        row_cols = st.columns([3, 1, 1, 1, 1, 1, 1, 1, 1, 1])
        question_text = row_cols[0].text_input(
            "Question Description",
            value=current_row["Question Description"],
            key=f"question_description_{index}",
            label_visibility="collapsed",
        )
        s_score = row_cols[1].number_input(
            "S",
            min_value=0.0,
            max_value=1.0,
            step=0.05,
            value=float(current_row["S"]),
            key=f"s_score_{index}",
            label_visibility="collapsed",
        )
        t_score = row_cols[2].number_input(
            "T/E",
            min_value=0.0,
            max_value=1.0,
            step=0.05,
            value=float(current_row["T/E"]),
            key=f"t_score_{index}",
            label_visibility="collapsed",
        )
        a_score = row_cols[3].number_input(
            "A",
            min_value=0.0,
            max_value=1.0,
            step=0.05,
            value=float(current_row["A"]),
            key=f"a_score_{index}",
            label_visibility="collapsed",
        )
        r_score = row_cols[4].number_input(
            "R/T",
            min_value=0.0,
            max_value=1.0,
            step=0.05,
            value=float(current_row["R/T"]),
            key=f"r_score_{index}",
            label_visibility="collapsed",
        )
        star_rating = round((s_score + t_score + a_score + r_score) / 4, 2)
        with row_cols[5]:
            render_readonly_score(star_rating)
        c_score = row_cols[6].number_input(
            "C/S",
            min_value=0.0,
            max_value=1.0,
            step=0.05,
            value=float(current_row["C/S"]),
            key=f"c_score_{index}",
            label_visibility="collapsed",
        )
        ae_score = row_cols[7].number_input(
            "A/E",
            min_value=0.0,
            max_value=1.0,
            step=0.05,
            value=float(current_row["A/E"]),
            key=f"ae_score_{index}",
            label_visibility="collapsed",
        )
        ra_score = row_cols[8].number_input(
            "R/A",
            min_value=0.0,
            max_value=1.0,
            step=0.05,
            value=float(current_row["R/A"]),
            key=f"ra_score_{index}",
            label_visibility="collapsed",
        )
        competency_rating = round((c_score + ae_score + ra_score) / 3, 2)
        with row_cols[9]:
            render_readonly_score(competency_rating)
        updated_rows.append(
            {
                "Question Description": question_text,
                "S": s_score,
                "T/E": t_score,
                "A": a_score,
                "R/T": r_score,
                "STAR Rating": star_rating,
                "C/S": c_score,
                "A/E": ae_score,
                "R/A": ra_score,
                "Competency Rating": competency_rating,
            }
        )

    scored_rows = recalculate_score_rows(updated_rows)
    st.session_state["score_rows"] = scored_rows
    return scored_rows


def render_metric_section(scored_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    st.markdown("### Metric")
    metric_rows = build_metric_rows(scored_rows)
    render_simple_table(metric_rows, ["Metric", "Value"])
    return metric_rows


def render_interview_report_section(
    scored_rows: list[dict[str, Any]], metric_rows: list[dict[str, Any]]
) -> None:
    st.markdown("### Report")
    prompt_file = st.file_uploader(
        "Upload Evaluation Prompt",
        type=["doc", "docx", "pdf", "txt"],
        key="evaluation_prompt",
    )

    if st.button("Generate Interview Evaluation Report", use_container_width=True):
        try:
            prompt_text, prompt_source = resolve_prompt_text(prompt_file, PROMPT_KEYS["interview_analysis"])
            if not prompt_text:
                st.error("Please upload a prompt document or ask an admin to configure the default interview analysis prompt.")
                return
            if prompt_source == "admin_default":
                st.info("Using the admin default interview analysis prompt.")

            payload = {
                "task": "interview_evaluation_report",
                "scoring_table": scored_rows,
                "metric_table": metric_rows,
                "question_simulation_report": st.session_state.get("question_report_text", ""),
                "custom_prompt": prompt_text,
                **build_reference_payload(),
            }

            with st.spinner("Generating interview evaluation report..."):
                report_body = call_llm(
                    system_prompt=(
                        "You are an interview performance analyst. Follow the uploaded prompt carefully "
                        "and produce a structured English evaluation report."
                    ),
                    model=st.session_state["selected_model"],
                    available_models=st.session_state["available_models"],
                    payload=payload,
                )

            persist_report(
                "interview_evaluation_report",
                "interview_evaluation_report.docx",
                "Interview Evaluation Report",
                report_body,
                [
                    ("Evaluation Report", report_body),
                    ("Scoring Table", json.dumps(scored_rows, ensure_ascii=False, indent=2)),
                    ("Metric Table", json.dumps(metric_rows, ensure_ascii=False, indent=2)),
                ],
            )

            st.session_state["evaluation_report_text"] = report_body
            st.success("Interview evaluation report generated.")
        except Exception as exc:
            st.error(str(exc))

    render_persisted_report(
        "interview_evaluation_report",
        "Interview Evaluation Report",
        "Download Interview Evaluation Report",
    )


def clear_admin_login_inputs() -> None:
    st.session_state["clear_admin_login_inputs"] = True


def apply_pending_admin_input_clear() -> None:
    if st.session_state.get("clear_admin_login_inputs"):
        st.session_state.pop("admin_username_input", None)
        st.session_state.pop("admin_password_input", None)
        st.session_state["clear_admin_login_inputs"] = False


@st.dialog("Admin Login")
def render_admin_login_dialog() -> None:
    st.write("Enter the administrator credentials.")
    username = st.text_input("Username", key="admin_username_input")
    password = st.text_input("Password", type="password", key="admin_password_input")

    login_col, cancel_col = st.columns(2)
    with login_col:
        if st.button("Log In", use_container_width=True):
            if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
                st.session_state["is_admin"] = True
                st.session_state["show_admin_login_dialog"] = False
                st.session_state["admin_login_error"] = ""
                clear_admin_login_inputs()
                if st.session_state.get("api_key_input") == ADMIN_PASSWORD:
                    st.session_state["api_key_input"] = ""
                st.rerun()
            else:
                st.session_state["admin_login_error"] = "Invalid administrator username or password."
                st.rerun()
    with cancel_col:
        if st.button("Cancel", use_container_width=True):
            st.session_state["show_admin_login_dialog"] = False
            st.session_state["admin_login_error"] = ""
            clear_admin_login_inputs()
            st.rerun()

    if st.session_state.get("admin_login_error"):
        st.error(st.session_state["admin_login_error"])


def render_manage_tab() -> None:
    st.subheader("Manage")
    st.caption("Upload the default prompt files used when regular users leave prompt upload empty.")

    prompt_sections = [
        ("Resume Default Prompt", PROMPT_KEYS["resume"], "manage_resume_prompt"),
        ("Interview Question Default Prompt", PROMPT_KEYS["interview_questions"], "manage_interview_questions_prompt"),
        ("Interview Analysis Default Prompt", PROMPT_KEYS["interview_analysis"], "manage_interview_analysis_prompt"),
    ]

    config = read_admin_prompt_config()
    for label, prompt_key, uploader_key in prompt_sections:
        st.markdown(f"### {label}")
        st.write(f"Current file: `{config.get(prompt_key, 'Not configured')}`")
        uploaded_file = st.file_uploader(
            f"Upload {label}",
            type=["doc", "docx", "pdf", "txt"],
            key=uploader_key,
        )
        if uploaded_file is not None:
            upload_signature = f"{uploaded_file.name}:{len(uploaded_file.getvalue())}"
            signature_key = f"{prompt_key}_saved_signature"
            if st.session_state.get(signature_key) != upload_signature:
                saved_path = save_admin_prompt(uploaded_file, prompt_key)
                st.session_state[signature_key] = upload_signature
                config[prompt_key] = saved_path.name
                st.success(f"Saved default prompt to `{saved_path.name}`.")

    st.markdown("### Reference Libraries")
    st.write(f"`anzsco` folder: `{ANZSCO_DIR}`")
    st.write(f"`sifa` folder: `{SIFA_DIR}`")
    st.write("Supported reference file types: `pdf`, `doc`, `docx`, `csv`, `xlsx`, `txt`.")


def main() -> None:
    st.set_page_config(page_title="Career Copilot MVP", page_icon=":memo:", layout="wide")
    ensure_app_directories()

    if "api_key_input" not in st.session_state:
        st.session_state["api_key_input"] = os.getenv("OPENROUTER_API_KEY") or os.getenv("OPENAI_API_KEY") or ""
    if "is_admin" not in st.session_state:
        st.session_state["is_admin"] = False
    if "show_admin_login_dialog" not in st.session_state:
        st.session_state["show_admin_login_dialog"] = False
    if "admin_login_error" not in st.session_state:
        st.session_state["admin_login_error"] = ""
    if "clear_admin_login_inputs" not in st.session_state:
        st.session_state["clear_admin_login_inputs"] = False

    apply_pending_admin_input_clear()

    available_models, model_fetch_warning = get_available_models()
    st.session_state["available_models"] = available_models
    available_model_ids = [model["id"] for model in available_models]
    if "selected_model" not in st.session_state or st.session_state["selected_model"] not in available_model_ids:
        st.session_state["selected_model"] = available_models[0]["id"]

    st.title("Career Copilot MVP")
    st.caption("Upload your prompts and documents, then generate Word reports in English.")

    with st.sidebar:
        st.header("Configuration")
        st.session_state["selected_model"] = st.selectbox(
            "Model",
            options=available_model_ids,
            index=available_model_ids.index(st.session_state["selected_model"]),
            format_func=lambda model_id: get_model_by_id(available_models, model_id)["label"],
            key="selected_model_widget",
        )
        st.text_input(
            "API Key",
            key="api_key_input",
            type="password",
            help="Paste your OpenRouter API key here. This value is kept only in the current session.",
        )
        selected_model = get_model_by_id(available_models, st.session_state["selected_model"])
        pricing = selected_model.get("pricing", {})
        st.caption(
            "Selected model details: "
            f"Input {format_price_per_million(pricing.get('prompt'))}, "
            f"Output {format_price_per_million(pricing.get('completion'))}, "
            f"Context {format_context_length(selected_model.get('context_length'))}"
        )
        if st.session_state.get("api_key_input") or os.getenv("OPENROUTER_API_KEY"):
            st.write("Provider: `OpenRouter`")
            st.write("Models are fetched dynamically and filtered to `:free` only.")
            if model_fetch_warning:
                st.warning(model_fetch_warning)
        else:
            st.write("Provider: `OpenAI-compatible default`")
            st.write("Set `OPENROUTER_API_KEY` to enable dynamic OpenRouter free-model filtering.")
        st.write("Set `OPENROUTER_API_KEY` or `OPENAI_API_KEY` before generating reports.")
        st.write("Supported uploads: `DOC`, `DOCX`, `PDF`, `TXT`, `CSV`, `XLSX`.")
        st.write(f"Current selected model: `{st.session_state['selected_model']}`")
        if st.session_state.get("active_request_model"):
            st.write(f"Last request model: `{st.session_state['active_request_model']}`")
        st.write("---")
        if st.session_state["is_admin"]:
            st.success("Logged in as admin.")
            if st.button("Log Out", use_container_width=True):
                st.session_state["is_admin"] = False
                st.session_state["show_admin_login_dialog"] = False
                st.session_state["admin_login_error"] = ""
                clear_admin_login_inputs()
                st.rerun()
        else:
            if st.button("Log In", use_container_width=True):
                st.session_state["show_admin_login_dialog"] = True

    if st.session_state.get("show_admin_login_dialog"):
        render_admin_login_dialog()

    if st.session_state["is_admin"]:
        resume_tab, interview_tab, manage_tab = st.tabs(["Resume", "Interview", "Manage"])
    else:
        resume_tab, interview_tab = st.tabs(["Resume", "Interview"])

    with resume_tab:
        render_resume_module()

    with interview_tab:
        render_question_simulation()
        scored_rows = render_scoring_section()
        metric_rows = render_metric_section(scored_rows)
        render_interview_report_section(scored_rows, metric_rows)

    if st.session_state["is_admin"]:
        with manage_tab:
            render_manage_tab()


if __name__ == "__main__":
    main()
